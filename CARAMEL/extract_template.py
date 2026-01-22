import json
from docx import Document
import os

def extract_features(docx_path, output_json):
    if not os.path.exists(docx_path):
        print(f"Error: {docx_path} not found.")
        return

    doc = Document(docx_path)
    features = {
        "styles": [],
        "structure": []
    }

    # Extract Standard Styles used (heuristic)
    # python-docx doesn't easily list *all* available styles in the doc directly without iterating, 
    # but we can see what's used.
    used_styles = set()

    for para in doc.paragraphs:
        style_name = para.style.name
        used_styles.add(style_name)
        
        item = {
            "type": "paragraph",
            "style": style_name,
            "text": para.text.strip()
        }
        if item["text"]: # Only save non-empty structure for clarity
            features["structure"].append(item)

    # Tables
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        
        features["structure"].append({
            "type": "table",
            "content": table_data
        })

    features["styles"] = list(used_styles)

    with open(output_json, 'w', encoding='utf-8') as f:
        json.dump(features, f, indent=4)
        
    print(f"Features extracted to {output_json}")

if __name__ == "__main__":
    extract_features("SOW_TEMPLATE.docx", "template_features.json")
