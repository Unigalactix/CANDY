import os
from openai import OpenAI, AzureOpenAI
from pypdf import PdfReader
from xhtml2pdf import pisa
from io import BytesIO

def extract_text_from_pdf(pdf_path):
    """
    Extracts text content from a PDF file.
    """
    try:
        reader = PdfReader(pdf_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        return f"Error reading PDF template: {str(e)}"

from pypdf import PdfReader, PdfWriter, PageObject, Transformation
from docx import Document

def generate_merged_pdf(html_content, template_path):
    """
    Generates a PDF from HTML and overlays it onto the template PDF.
    Preserves formatting/logos of the template.
    """
    # 1. Generate Content PDF (Transparent Background)
    # Inject CSS for transparency if not present
    if "@page" not in html_content:
        html_content = f"""
        <style>
            @page {{
                background-color: transparent;
                margin-top: 4cm; /* Adjust based on template header size */
                margin-bottom: 2cm; /* Adjust based on template footer size */
                margin-left: 2cm;
                margin-right: 2cm;
            }}
        </style>
        {html_content}
        """
    
    content_pdf_buffer = BytesIO()
    pisa.CreatePDF(BytesIO(html_content.encode('utf-8')), content_pdf_buffer)
    content_pdf_buffer.seek(0)
    
    # 2. Merge with Template
    try:
        template_reader = PdfReader(template_path)
        content_reader = PdfReader(content_pdf_buffer)
        writer = PdfWriter()
        
        # Iterate through content pages
        for i, content_page in enumerate(content_reader.pages):
            # Get corresponding template page (loop if content > template)
            template_page_idx = i % len(template_reader.pages)
            template_page = template_reader.pages[template_page_idx]
            
            # Create a blank page with template dimensions
            output_page = PageObject.create_blank_page(
                width=template_page.mediabox.width,
                height=template_page.mediabox.height
            )
            
            # Merge Template (Background)
            output_page.merge_page(template_page)
            
            # Merge Content (Foreground)
            output_page.merge_page(content_page)
            
            writer.add_page(output_page)
            
        result = BytesIO()
        writer.write(result)
        return result.getvalue()
        
    except Exception as e:
        print(f"Error merging PDF: {e}")
        # Fallback to simple PDF if merge fails
        return content_pdf_buffer.getvalue()

from pypdf import PdfReader, PdfWriter, PageObject, Transformation
from docx import Document
import tempfile
import pythoncom
from docx2pdf import convert

def generate_merged_pdf(html_content, template_path):
    # ... (Overlay Logic - Keeping this if needed for fallback, but user prefers DOCX template) ...
    # Simplified for readability in this context, assuming we rely on DOCX conversion now.
    pass

def generate_sow_doc_struct(sow_content, template_path):
    """
    Appends the SOW content to the DOCX template and returns the Document object.
    """
    try:
        doc = Document(template_path)
        doc.add_page_break()
        
        lines = sow_content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
                
            if line.startswith('## '):
                doc.add_heading(line.replace('## ', ''), level=2)
            elif line.startswith('# '):
                doc.add_heading(line.replace('# ', ''), level=1)
            elif line.startswith('- '):
                # Try 'List Bullet' first, fallback to 'List Paragraph', then plain
                text = line.replace('- ', '')
                try:
                    doc.add_paragraph(text, style='List Bullet')
                except:
                    try:
                        doc.add_paragraph(text, style='List Paragraph')
                    except:
                        doc.add_paragraph(f"• {text}") # Manual bullet
            else:
                doc.add_paragraph(line)
        return doc
    except Exception as e:
        raise ValueError(f"Error structuring DOCX: {e}")

def generate_sow_docx_bytes(sow_content, template_path):
    """Generates DOCX bytes."""
    doc = generate_sow_doc_struct(sow_content, template_path)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

def convert_docx_to_pdf_bytes(doc_bytes):
    """
    Converts DOCX bytes to PDF bytes using docx2pdf (requires Word).
    """
    try:
        # Initialize COM for Streamlit thread
        pythoncom.CoInitialize()
        
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:
            tmp_docx.write(doc_bytes)
            docx_path = tmp_docx.name
            
        pdf_path = docx_path.replace(".docx", ".pdf")
        
        try:
            convert(docx_path, pdf_path)
            with open(pdf_path, "rb") as f:
                pdf_bytes = f.read()
            return pdf_bytes
        finally:
            # Cleanup
            if os.path.exists(docx_path):
                os.remove(docx_path)
            if os.path.exists(pdf_path):
                os.remove(pdf_path)
                
    except Exception as e:
        raise ValueError(f"PDF Conversion Failed (Ensure Word is installed): {e}")
    finally:
        pythoncom.CoUninitialize()


def get_llm_client():
    """
    Returns the appropriate OpenAI or AzureOpenAI client based on .env configuration.
    """
    azure_endpoint = os.getenv("AZURE_OPENAI_ENDPOINT")
    azure_api_key = os.getenv("AZURE_OPENAI_API_KEY")
    azure_api_version = os.getenv("AZURE_OPENAI_API_VERSION", "2024-02-15-preview")
    
    if azure_endpoint and azure_api_key:
        print(f"[+] Connecting to Azure OpenAI at {azure_endpoint}...")
        return AzureOpenAI(
            azure_endpoint=azure_endpoint,
            api_key=azure_api_key,
            api_version=azure_api_version
        ), os.getenv("AZURE_OPENAI_DEPLOYMENT")
    
    base_url = os.getenv("LLM_BASE_URL")
    api_key = os.getenv("LLM_API_KEY")
    model = os.getenv("LLM_MODEL")
    
    if base_url and api_key and model:
        print(f"[+] Connecting to Local/Standard LLM at {base_url}...")
        return OpenAI(
            base_url=base_url,
            api_key=api_key
        ), model
        
    raise ValueError("Missing LLM configuration. Check .env variables.")

def generate_sow_draft(mom_text):
    """
    Generates a text/markdown draft of the SOW from MOM details.
    Returns: Markdown text string.
    """
    try:
        client, model = get_llm_client()
    except ValueError as e:
        return str(e)

    # --- Processing Logic for Large Inputs ---
    
    # helper: Chunk text
    def chunk_text(text, size=2000):
        return [text[i:i+size] for i in range(0, len(text), size)]
    
    mom_chunks = chunk_text(mom_text)
    consolidated_info = ""
    
    if len(mom_chunks) > 1:
        print(f"[+] Processing MOM in {len(mom_chunks)} batches...")
        for i, chunk in enumerate(mom_chunks, 1):
            print(f"    - Processing Batch {i}/{len(mom_chunks)}...")
            summary_prompt = (
                f"I have a section of Meeting Minutes. Extract key project details (Scope, Timeline, Budget, Team, Deliverables, etc.) as concise bullet points.\n"
                f"Ignore conversational filler.\n\n"
                f"MOM Segment:\n"
                f"{chunk}"
            )
            try:
                resp = client.chat.completions.create(
                    model=model,
                    messages=[{"role": "user", "content": summary_prompt}],
                    temperature=0.3
                )
                consolidated_info += f"\n--- Batch {i} Details ---\n{resp.choices[0].message.content}\n"
            except Exception as e:
                consolidated_info += f"\n[Error extracting Batch {i}: {str(e)}]\n"
    else:
        consolidated_info = mom_text

    # --- Draft Generation ---
    print("[+] Generating SOW Draft...")
    
    system_prompt = "You are a professional Project Manager and Technical Writer."
    user_prompt = (
        f"I have the following Consolidated Project Details (extracted from meeting minutes):\n"
        f"---------------------\n"
        f"{consolidated_info}\n"
        f"---------------------\n\n"
        f"Instructions:\n"
        f"1. Create a detailed Statement of Work (SOW) draft.\n"
        f"2. Use Markdown formatting (## Headers, - Bullet points).\n"
        f"3. Include standard SOW sections: Project Overview, Scope of Work, Deliverables, Timeline, Pricing/Budget, Governance/Team.\n"
        f"4. Do NOT use HTML tags yet. Just structure the content clearly.\n"
    )
    
    try:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.3,
        )
        content = response.choices[0].message.content
        return content
        
    except Exception as e:
        return f"Error Generating SOW Draft: {str(e)}"

def format_sow_to_html(edited_text, template_pdf_path):
    """
    Takes the edited SOW text and wraps it in the HTML structure of the template PDF.
    """
    try:
        client, model = get_llm_client()
    except ValueError as e:
        return f"<h3>Error: Missing Configuration</h3><p>{str(e)}</p>"
    
    # 1. Extract Template Text
    print(f"\n[+] Reading Template from {template_pdf_path}...")
    template_text = extract_text_from_pdf(template_pdf_path)
    
    # 2. Generate HTML
    print("[+] Applying Template Formatting...")
    
    system_prompt = "You are a specialized document formatter."
    user_prompt = (
        f"I have the final SOW Content (Markdown):\n"
        f"---------------------\n"
        f"{edited_text}\n"
        f"---------------------\n\n"
        f"And the Style/Structure extracted from a Reference PDF via OCR/Text Extraction:\n"
        f"---------------------\n"
        f"{template_text}\n"
        f"---------------------\n\n"
        f"Instructions:\n"
        f"1. Convert the 'SOW Content' into an HTML document.\n"
        f"2. Mimic the structure and specific standard clauses found in the 'Reference PDF' where applicable, but keep the specific project details from 'SOW Content'.\n"
        f"3. Use HTML tags (<h1>, <p>, <ul> etc.).\n"
        f"4. Output ONLY the valid HTML. Do not include markdown code blocks.\n"
    )
    
    try:
        response = client.chat.completions.create(
            model=model,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.1, # Low temp for strict formatting
        )
        content = response.choices[0].message.content
        content = content.replace("```html", "").replace("```", "").strip()
        return content
    except Exception as e:
        return f"<h3>Error Formatting SOW</h3><p>{str(e)}</p>"
